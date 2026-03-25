"""
Resume Formatter — fills NexTurn Word template by direct w:t element replacement.

"""

import io
import os
import json
import copy
import sys
import traceback
import re

import streamlit as st

from backend.openai_client import create_openai_completion
from backend.preprocessing import parse_resume_locally


def _import_docx_document():
    try:
        from docx import Document
    except Exception:
        try:
            from docx.api import Document
        except Exception as e_api:
            raise ImportError(
                "Could not import Document from python-docx. "
                "Uninstall the legacy 'docx' package and ensure 'python-docx' is installed."
            ) from e_api
    from docx.oxml.ns import qn
    return Document, qn


Document, qn = _import_docx_document()

# Template path
TEMPLATE_PATH = os.path.join(
    os.path.dirname(os.path.dirname(os.path.abspath(__file__))),
    "inputs", "word_template.docx"
)

# LLM extraction

def extract_detailed_resume_data(client, resume_text: str, candidate_meta: dict) -> dict:
    """
    Extract detailed resume data for Word-template generation.

    KEY FIX: Now returns BOTH the flat keys (for Word template) AND
    a key_projects list (for PPT template_mapper.py), so both downstream
    generators always get the data they need without "Information Missing".
    """

    prompt = f"""You are an expert resume parser. Extract ALL details from the resume below into structured JSON.

RULES:
- Extract exact text from the resume. Do NOT fabricate any information.
- For bullet points, extract the actual bullet text verbatim (no leading dashes, bullets, or numbers).
- Every field that is not found must be an empty string "", not null or "N/A".
- responsibilities in key_projects must always be a LIST of strings.
- Return ONLY this JSON (no markdown, no explanation):

{{
  "NAME": "Full name of the candidate",
  "ROLE": "Current or most recent job title",
  "PROFESSIONAL_SUMMARY": "2-3 sentence professional summary",
  "experience_years": "Total years as a number e.g. 5.5",
  "COMPANY_NAME": "Most recent company name",
  "LOCATION": "City, Country",
  "START_DATE": "Month YYYY of most recent role start",
  "END_DATE": "Month YYYY or Present",

  "PROJECT1_NAME": "Primary project/role title",
  "ABOUT_PROJECT_BULLET_1": "First responsibility bullet",
  "ABOUT_PROJECT_BULLET_2": "Second responsibility bullet",
  "ABOUT_PROJECT_BULLET_3": "Third responsibility bullet",
  "ABOUT_PROJECT_BULLET_4": "Fourth responsibility bullet",
  "ABOUT_PROJECT_BULLET_5": "Fifth responsibility bullet",
  "ABOUT_PROJECT_BULLET_6": "Sixth responsibility bullet",

  "PROJECT2_NAME": "Second project/role title or empty string",
  "PROJECT2_BULLET_1": "First responsibility bullet",
  "PROJECT2_BULLET_2": "Second responsibility bullet",
  "PROJECT2_BULLET_3": "Third responsibility bullet",
  "PROJECT2_BULLET_4": "Fourth responsibility bullet",
  "PROJECT2_BULLET_5": "Fifth responsibility bullet",

  "PROJECT3_NAME": "Third project/role title or empty string",
  "PROJECT3_BULLET_1": "First responsibility bullet",
  "PROJECT3_BULLET_2": "Second responsibility bullet",
  "PROJECT3_BULLET_3": "Third responsibility bullet",
  "PROJECT3_BULLET_4": "Fourth responsibility bullet",
  "PROJECT3_BULLET_5": "Fifth responsibility bullet",

  "PROJECT4_NAME": "Fourth project/role title or empty string",
  "PROJECT4_BULLET_1": "First responsibility bullet",
  "PROJECT4_BULLET_2": "Second responsibility bullet",
  "PROJECT4_BULLET_3": "Third responsibility bullet",
  "PROJECT4_BULLET_4": "Fourth responsibility bullet",
  "PROJECT4_BULLET_5": "Fifth responsibility bullet",

  "TECHNOLOGIES_USED": "Comma separated list of all technologies used",
  "HIGHEST_EDUCATION": "Full degree e.g. B.E. in Computer Science Engineering",
  "COLLEGE_NAME": "University or college name",
  "EDUCATION_DATES": "Month YYYY – Month YYYY or YYYY – YYYY",
  "tech_stack": ["Skill1","Skill2","Skill3","Skill4","Skill5","Skill6","Skill7","Skill8","Skill9","Skill10","Skill11","Skill12"],
  "BACKEND_LANGUAGES": "e.g. Python, Java",
  "CONTAINERS_AND_ORCHESTRATION": "e.g. Docker, Kubernetes",
  "DATABASES": "e.g. PostgreSQL, MongoDB",
  "OPERATING_SYSTEMS": "e.g. Linux, Windows",
  "VERSION_CONTROL_TOOLS": "e.g. Git, GitHub",
  "TESTING_TOOLS": "e.g. PyTest, Selenium"
}}

Resume:
{resume_text[:7000]}"""

    # Local heuristic snapshot (no resume content logged)
    local = {}
    try:
        local = parse_resume_locally(resume_text or "")
    except Exception:
        local = {}

    data = None
    try:
        response = create_openai_completion(
            client,
            messages=[
                {"role": "system", "content": "You are a precise resume parser. Return only valid JSON."},
                {"role": "user",   "content": prompt},
            ],
            model="gpt-4o-mini",
            temperature=0.1,
            max_tokens=3000,
        )
        content = response.choices[0].message.content.strip()
        j_start = content.find("{")
        j_end   = content.rfind("}") + 1
        if j_start != -1 and j_end > j_start:
            data = json.loads(content[j_start:j_end])
    except Exception as e:
        st.warning(f"Could not fully parse resume details: {e}")

    if not data:
        data = {}

    # Fill top-level fields from candidate_meta if AI missed them 
    def _empty(val):
        return not val or str(val).strip() in ("", "null", "None", "N/A", "n/a", "nan")

    if _empty(data.get("NAME")):
        data["NAME"] = candidate_meta.get("name", "")
    if _empty(data.get("ROLE")):
        data["ROLE"] = candidate_meta.get("current_role", "")
    if _empty(data.get("experience_years")):
        data["experience_years"] = str(candidate_meta.get("experience_years", ""))
    if _empty(data.get("PROFESSIONAL_SUMMARY")):
        data["PROFESSIONAL_SUMMARY"] = candidate_meta.get("objective", "")
    if _empty(data.get("HIGHEST_EDUCATION")):
        data["HIGHEST_EDUCATION"] = candidate_meta.get("education", "")
    if _empty(data.get("HIGHEST_EDUCATION")):
        data["HIGHEST_EDUCATION"] = local.get("education", "")

    # Ensure tech_stack is always a list 
    ts = data.get("tech_stack")
    if isinstance(ts, str):
        data["tech_stack"] = [s.strip() for s in ts.split(",") if s.strip()]
    elif not isinstance(ts, list):
        # Try pulling from candidate_meta
        meta_ts = candidate_meta.get("tech_stack", "")
        if isinstance(meta_ts, str):
            data["tech_stack"] = [s.strip() for s in meta_ts.split(",") if s.strip()]
        elif isinstance(meta_ts, list):
            data["tech_stack"] = [str(s).strip() for s in meta_ts if s]
        else:
            data["tech_stack"] = []

    # If TECHNOLOGIES_USED is empty, derive from tech_stack 
    if _empty(data.get("TECHNOLOGIES_USED")) and data.get("tech_stack"):
        data["TECHNOLOGIES_USED"] = ", ".join(str(s) for s in data["tech_stack"][:20])

    # Build key_projects list from flat PROJECT{n} keys 
    # This is the critical fix: template_mapper.py prefers key_projects list.
    # We build it here from the flat keys so BOTH Word and PPT generators work.
    key_projects = []
    for i in range(1, 5):
        if i == 1:
            name_key   = "PROJECT1_NAME"
            bullet_prefix = "ABOUT_PROJECT_BULLET_"
            bullet_range  = range(1, 7)
        else:
            name_key      = f"PROJECT{i}_NAME"
            bullet_prefix = f"PROJECT{i}_BULLET_"
            bullet_range  = range(1, 6)

        proj_name = str(data.get(name_key, "") or "").strip()
        bullets   = []
        for b in bullet_range:
            bval = str(data.get(f"{bullet_prefix}{b}", "") or "").strip()
            if bval:
                bullets.append(bval)

        # Try to split duration/role from PROJECT{i}_NAME if it contains " - " or ": "
        duration = ""
        role_str = ""
        if i > 1:
            duration = str(data.get(f"PROJECT{i}_DURATION", "") or "").strip()
            role_str = str(data.get(f"PROJECT{i}_ROLE", "") or "").strip()

        if i == 1:
            duration = f"{data.get('START_DATE', '')} – {data.get('END_DATE', '')}".strip(" –")
            role_str = data.get("ROLE", "")

        key_projects.append({
            "title":            proj_name,
            "duration":         duration,
            "role":             role_str,
            "description":      "",   # Word template doesn't use this; PPT uses it optionally
            "responsibilities": bullets,
        })

    # Pad to 4
    while len(key_projects) < 4:
        key_projects.append({
            "title": "", "duration": "", "role": "",
            "description": "", "responsibilities": [],
        })

    data["key_projects"] = key_projects

    return data

# Template completeness check

def check_template_completeness(data: dict) -> dict:
    warnings = []
    if not data.get("NAME"):               warnings.append("Candidate name not found")
    if not data.get("ROLE"):               warnings.append("Job title not found")
    if not data.get("experience_years"):   warnings.append("Years of experience not specified")
    if not data.get("COMPANY_NAME"):       warnings.append("Company name missing")
    if not data.get("PROJECT1_NAME"):      warnings.append("Primary project not found")
    if not data.get("PROFESSIONAL_SUMMARY"): warnings.append("Professional summary missing")
    if not data.get("HIGHEST_EDUCATION"):  warnings.append("Education details not found")
    critical = not data.get("NAME") or not data.get("COMPANY_NAME")
    return {"complete": len(warnings) == 0, "warnings": warnings, "has_critical_gaps": critical}

# Paragraph helpers

def _clear_and_set_para(para, new_text, bold=None):
    """
    Replace text runs in para with new_text, preserving drawing/anchor elements.
    """
    p_elem = para._p
    runs = p_elem.findall(qn("w:r"))
    if not runs:
        para.add_run(new_text)
        return

    MC_NS = "http://schemas.openxmlformats.org/markup-compatibility/2006"

    def _is_drawing_run(r):
        if r.find(qn("w:drawing")) is not None: return True
        if r.find(qn("w:pict"))    is not None: return True
        if r.find(f"{{{MC_NS}}}AlternateContent") is not None: return True
        return False

    text_runs    = [r for r in runs if not _is_drawing_run(r)]
    drawing_runs = [r for r in runs if     _is_drawing_run(r)]

    ref_run  = text_runs[0] if text_runs else runs[0]
    rpr      = ref_run.find(qn("w:rPr"))
    rpr_copy = copy.deepcopy(rpr) if rpr is not None else None

    if bold is not None and rpr_copy is not None:
        b_elem = rpr_copy.find(qn("w:b"))
        if bold and b_elem is None:
            rpr_copy.insert(0, _make_elem("w:b"))
        elif not bold and b_elem is not None:
            rpr_copy.remove(b_elem)

    for r in text_runs:
        p_elem.remove(r)

    new_r = _make_elem("w:r")
    if rpr_copy is not None:
        new_r.append(rpr_copy)
    t = _make_elem("w:t")
    t.text = new_text
    if new_text and (new_text[0] == " " or new_text[-1] == " "):
        t.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
    new_r.append(t)

    if drawing_runs:
        drawing_runs[0].addprevious(new_r)
    else:
        p_elem.append(new_r)


def _make_elem(tag):
    from docx.oxml import OxmlElement
    return OxmlElement(tag)


def _set_t(t_elem, text):
    t_elem.text = text
    if text and (text[0] == " " or text[-1] == " "):
        t_elem.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
    else:
        attr = "{http://www.w3.org/XML/1998/namespace}space"
        if attr in t_elem.attrib:
            del t_elem.attrib[attr]


def _get_anchor_skill_slots(body):
    anchors = body.findall("./" + qn("wp:anchor"))
    if not anchors:
        anchors = body.findall(".//" + qn("wp:anchor"))
    if not anchors:
        return []
    anchor0 = anchors[0]
    anchor0_t = anchor0.findall(".//" + qn("w:t"))
    skip = {"ACADEMIC", "QUALIFICATIONS", "KEY", "SKILLS", "OBJECTIVE", "•", ""}
    slots = []
    for t in anchor0_t:
        txt = (t.text or "").strip()
        if txt and txt not in skip:
            slots.append(t)
    return slots

# Main Word document generator

def generate_resume_docx(data: dict) -> bytes | None:
    try:
        if not os.path.exists(TEMPLATE_PATH):
            st.error(f"Word template not found at: {TEMPLATE_PATH}")
            return None

        buf_in = io.BytesIO()
        with open(TEMPLATE_PATH, "rb") as f:
            buf_in.write(f.read())
        buf_in.seek(0)
        doc  = Document(buf_in)
        body = doc.element.body

        def v(key, fallback=""):
            val = data.get(key, fallback)
            return str(val).strip() if val else fallback

        name    = v("NAME",  "Candidate Name")
        role    = v("ROLE",  "Professional")
        exp_yrs = v("experience_years", "")
        summary = v("PROFESSIONAL_SUMMARY", "")
        company = v("COMPANY_NAME", "")
        location= v("LOCATION", "")
        start   = v("START_DATE", "")
        end     = v("END_DATE", "")
        p1_name = v("PROJECT1_NAME", "")
        p2_name = v("PROJECT2_NAME", "")
        tech    = v("TECHNOLOGIES_USED", "")
        edu_deg = v("HIGHEST_EDUCATION", "")
        edu_col = v("COLLEGE_NAME", "")
        edu_dat = v("EDUCATION_DATES", "")

        bullets1 = [v(f"ABOUT_PROJECT_BULLET_{i}") for i in range(1, 7)]
        bullets1 = [b for b in bullets1 if b]

        bullets2 = [v(f"PROJECT2_BULLET_{i}") for i in range(1, 6)]
        bullets2 = [b for b in bullets2 if b]

        tech_stack = data.get("tech_stack") or []
        if isinstance(tech_stack, str):
            tech_stack = [s.strip() for s in tech_stack.split(",") if s.strip()]
        tech_stack = [str(s) for s in tech_stack if s]

        paras = doc.paragraphs

        # [0] Name
        name_runs = [
            r for r in paras[0]._p.findall(qn("w:r"))
            if r.find(qn("w:t")) is not None
            and (r.find(qn("w:t")).text or "").strip()
            and (r.find(qn("w:t")).text or "").strip() not in ("\t", " ")
        ]
        if name_runs:
            name_runs[0].find(qn("w:t")).text = name.upper()
        else:
            _clear_and_set_para(paras[0], f"\t{name.upper()}\t ")

        # [1] Role
        MC_NS_R = "http://schemas.openxmlformats.org/markup-compatibility/2006"
        role_text_runs = [
            r for r in paras[1]._p.findall(qn("w:r"))
            if r.find(f"{{{MC_NS_R}}}AlternateContent") is None
            and r.find(qn("w:drawing")) is None
            and r.find(qn("w:t")) is not None
            and (r.find(qn("w:t")).text or "").strip() not in ("", "\t")
        ]
        if role_text_runs:
            t0 = role_text_runs[0].find(qn("w:t"))
            t0.text = role.upper()
            t0.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
            for r in role_text_runs[1:]:
                t = r.find(qn("w:t"))
                if t is not None:
                    t.text = " "
        else:
            _clear_and_set_para(paras[1], f"\t{role.upper()}\t ")

        # [3] Summary
        _clear_and_set_para(paras[3], summary)

        # [5] Work Experience heading
        exp_label = f"{exp_yrs}YRS" if exp_yrs else "N/A"
        _clear_and_set_para(paras[5], f"WORK EXPERIENCE ({exp_label}) ", bold=True)

        # Table: Company + Location
        if doc.tables:
            cell = doc.tables[0].cell(0, 0)
            for tp in cell.paragraphs:
                if tp.text.strip():
                    loc_str = f": {location}" if location else ""
                    _clear_and_set_para(tp, f"{company}{loc_str}", bold=True)
                    break

        # [8] Date range
        date_str = f"{start} – {end}" if start and end else start or end or ""
        _clear_and_set_para(paras[8], f"{date_str} ")

        # [9] Project 1 name
        _clear_and_set_para(paras[9], f"Project: {p1_name} " if p1_name else "Project: N/A ")

        # [10–15] Project 1 bullets
        for slot, para_idx in enumerate(range(10, 16)):
            if para_idx < len(paras):
                text = bullets1[slot] if slot < len(bullets1) else ""
                _clear_and_set_para(paras[para_idx], text)

        # [16] Project 2 name
        if len(paras) > 16:
            _clear_and_set_para(paras[16], f"Projects: {p2_name}" if p2_name else "", bold=True)

        # [17–21] Project 2 bullets
        for slot, para_idx in enumerate(range(17, 22)):
            if para_idx < len(paras):
                text = bullets2[slot] if slot < len(bullets2) else ""
                _clear_and_set_para(paras[para_idx], text)

        # [23] Technologies
        if len(paras) > 23:
            _clear_and_set_para(paras[23], f"Technologies: {tech}" if tech else "")

        # [26–37] SKILLS section
        skills_order = [
            ("BACKEND_LANGUAGES",           26, 27),
            ("CONTAINERS_AND_ORCHESTRATION", 28, 29),
            ("DATABASES",                   30, 31),
            ("OPERATING_SYSTEMS",           32, 33),
            ("VERSION_CONTROL_TOOLS",       34, 35),
            ("TESTING_TOOLS",               36, 37),
        ]
        for key, label_idx, val_idx in skills_order:
            if val_idx < len(paras):
                _clear_and_set_para(paras[val_idx], v(key) or "N/A")

        # Floating anchor: Education + Key Skills
        slots = _get_anchor_skill_slots(body)
        edu_assignments = {
            0: edu_deg,
            1: " ", 2: " ", 3: " ", 4: " ", 5: " ",
            6: edu_col,
            7: " ", 8: " ",
            9: edu_dat,
            10: " ", 11: " ",
        }
        for slot_idx, text in edu_assignments.items():
            if slot_idx < len(slots):
                _set_t(slots[slot_idx], text or " ")

        for i in range(16):
            slot_idx = 12 + i
            if slot_idx < len(slots):
                skill_text = tech_stack[i] if i < len(tech_stack) else " "
                _set_t(slots[slot_idx], skill_text or " ")

        buf_out = io.BytesIO()
        doc.save(buf_out)
        buf_out.seek(0)
        return buf_out.read()

    except Exception as e:
        st.error(f"Could not create document: {e}")
        import traceback
        st.error(traceback.format_exc())
        return None